import requests
import pandas as pd
from datetime import datetime, timedelta
import google.generativeai as genai
import os
from docx import Document

# 配置 Google Generative AI
genai.configure(api_key="key")

# 設定 API 端點和參數
base_url = "https://api.binance.com"
endpoint = "/api/v3/klines"
symbol = "ETHUSDT"  # ETH/USD 交易對
interval = "1d"  # 日線資料
limit = 90  # 過去 30 天

# 計算開始和結束時間
end_time = datetime.now()
start_time = end_time - timedelta(days=limit)

# 將時間轉換為毫秒
start_time_ms = int(start_time.timestamp() * 1000)
end_time_ms = int(end_time.timestamp() * 1000)

# 發送 API 請求
params = {
    "symbol": symbol,
    "interval": interval,
    "startTime": start_time_ms,
    "endTime": end_time_ms,
    "limit": limit
}

response = requests.get(base_url + endpoint, params=params)
data = response.json()

# 將資料轉換為 DataFrame
columns = ["Open time", "Open", "High", "Low", "Close", "Volume", "Close time", "Quote asset volume", "Number of trades", "Taker buy base asset volume", "Taker buy quote asset volume", "Ignore"]
df = pd.DataFrame(data, columns=columns)

# 計算日均價
df["Average Price"] = (df["High"].astype(float) + df["Low"].astype(float)) / 2

# 將時間戳轉換為可讀格式
df["Open time"] = pd.to_datetime(df["Open time"], unit='ms')

# 只保留需要的欄位
df = df[["Open time", "Average Price", "Volume"]]

# 將數據轉換為字符串格式，方便傳遞給 LLM
data_str = df.to_string()
print(data_str)

# 配置 LLM 模型
model = genai.GenerativeModel("gemini-1.5-flash")

# 生成內容
prompt = f"使用繁體中文回答,分析以下 ETH/USD 的歷史價格趨勢:\n{data_str}"
response = model.generate_content(prompt)

# 創建一個新的 Word 文件
doc = Document()
doc.add_heading('ETH/USD Historical Price Analysis', 0)

# 添加生成的內容到 Word 文件
doc.add_paragraph(response.text)

# 保存 Word 文件
doc.save('ETH_USD_Price_Analysis.docx')

print("內容已保存到 ETH_USD_Price_Analysis.docx")
