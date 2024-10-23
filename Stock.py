import yfinance as yf
import google.generativeai as genai
import os
from docx import Document

# 配置 Google Generative AI
genai.configure(api_key="key")

# 抓取台積電 (2330) 一年內的歷史價格
ticker = "2330.TW"
data = yf.download(ticker, period="1mo")

# 將數據轉換為字符串格式，方便傳遞給 LLM
data_str = data.to_string()
print(data_str)

# 配置 LLM 模型
model = genai.GenerativeModel("gemini-1.5-flash")

# 生成內容
prompt = f"使用繁體中文回答,分析以下公司歷史價格趨勢:\n{data_str}"
response = model.generate_content(prompt)

# 創建一個新的 Word 文件
doc = Document()
doc.add_heading('TSMC (2330) Historical Stock Price Analysis', 0)

# 添加生成的內容到 Word 文件
doc.add_paragraph(response.text)

# 保存 Word 文件
doc.save('TSMC_2330_Stock_Analysis.docx')

print("內容已保存到 TSMC_2330_Stock_Analysis.docx")
