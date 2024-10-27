import yfinance as yf
import google.generativeai as genai
from docx import Document
import pandas as pd

# 配置 Google Generative AI
genai.configure(api_key="key")

# 使用者輸入
ticker = input("請輸入股票代碼 (e.g., 2542.TW): ")
period = input("請輸入抓取年份 (e.g., 5y, 1y): ")
interval = input("請輸入抓取數據間隔 (e.g., 1mo, 1d): ")

# 抓取歷史價格數據
price_data = yf.download(ticker, period=period, interval=interval)
annual_price_data = price_data.resample('YE').mean()  # 每年的均價
monthly_price_data = price_data.resample('ME').mean()  # 每月均價
daily_price_data = price_data  # 每日均價

# 抓取股息、自由現金流和負債
dividends = yf.Ticker(ticker).dividends.resample('YE').sum()  # 每年的股息
financials = yf.Ticker(ticker).financials.transpose()
net_income = financials['Net Income']
total_debt = yf.Ticker(ticker).balance_sheet.loc['Total Debt']  # 負債總額
shares_outstanding = yf.Ticker(ticker).info['sharesOutstanding']
eps = net_income / shares_outstanding  # 每年的EPS
cash_dividends = yf.Ticker(ticker).cashflow.loc['Dividends Paid'] if 'Dividends Paid' in yf.Ticker(ticker).cashflow.index else 'N/A'  # 現金股利
stock_dividends = yf.Ticker(ticker).info['dividendYield'] * shares_outstanding if 'dividendYield' in yf.Ticker(ticker).info else 'N/A'  # 股票股利
free_cash_flow = yf.Ticker(ticker).cashflow.loc['Free Cash Flow'] if 'Free Cash Flow' in yf.Ticker(ticker).cashflow.index else 'N/A'  # 自由現金流

# 將數據轉換為字符串格式，方便傳遞給 LLM
data_str = (
    f"Annual Average Price:\n{annual_price_data['Close']}\n\n"
    f"Monthly Average Price:\n{monthly_price_data['Close']}\n\n"
    f"Daily Average Price:\n{daily_price_data['Close']}\n\n"
    f"Dividends:\n{dividends}\n\n"
    f"Cash Dividends:\n{cash_dividends}\n\n"
    f"Stock Dividends:\n{stock_dividends}\n\n"
    f"EPS:\n{eps}\n\n"
    f"Free Cash Flow:\n{free_cash_flow}\n\n"
    f"Total Debt:\n{total_debt}"
)
print(data_str)

# 配置 LLM 模型
model = genai.GenerativeModel("gemini-1.5-flash")

# 生成內容
prompt = (
    f"使用繁體中文回答, 分析以下公司的每年均價、股息、現金股利、股票股利、EPS、自由現金流和負債的趨勢, "
    f"並且假裝你是巴菲特給這股票評論優缺點和預測合理價:\n{data_str}"
)
response = model.generate_content(prompt)

# 創建分析報告 Word 文件
doc_analysis = Document()
doc_analysis.add_heading('Historical Stock Data Analysis', 0)

# 添加生成的內容到分析報告
doc_analysis.add_paragraph(response.text)

# 保存分析報告到指定路徑
doc_analysis.save('C:/Users/Desktop/Analysis.docx')

# 創建數據報告 Word 文件
doc_data = Document()
doc_data.add_heading('Historical Stock Data', 0)

# 添加數據到數據報告
doc_data.add_paragraph(data_str)

# 保存數據報告到指定路徑
doc_data.save('C:/Users/Desktop/Data.docx')

print("內容已保存")