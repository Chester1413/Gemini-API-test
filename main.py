
# 導入所需的模組
import fitz  # PyMuPDF
import google.generativeai as genai
from docx import Document

# 配置生成式 AI 模型
genai.configure(api_key="key")
model = genai.GenerativeModel("gemini-1.5-flash")
#對話記憶 會存進history
#chat = model.start_chat(history=[])
#response = chat.send_message("我想學習deep learning要從哪邊開始入手")

# 定義函數來提取 PDF 文件中的文本
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text

# 指定 PDF 文件的路徑
pdf_path = r"C:\Users\Mike\Desktop\GOOGLE\AAA.pdf"
pdf_text = extract_text_from_pdf(pdf_path)

# 使用生成式 AI 模型來生成中文重點
prompt = f"請使用繁體中文整理出財務報告的重點,最後並請你以Warren Edward Buffett的視角來評論這間公司的優點和缺點並做總結：\n\n{pdf_text}"
response = model.generate_content(prompt)

# 將生成的中文重點保存為 docx 文件
doc = Document()
doc.add_heading('PDF 中文重點', 0)
doc.add_paragraph(response.text)
doc.save('output.docx')

print("生成的中文重點已保存為 output.docx")
